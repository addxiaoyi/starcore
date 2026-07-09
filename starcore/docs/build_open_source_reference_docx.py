from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "STARCORE_Open_Source_Reference_Roadmap_2026-06-03.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
INK = "1F2937"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")
        tc_mar.append(tag)


def set_table_fixed(table, widths: list[float]) -> None:
    dxa_widths = [int(round(width * 1440)) for width in widths]
    if dxa_widths:
        dxa_widths[-1] += 9360 - sum(dxa_widths)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa_widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in dxa_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_paragraph(paragraph, size=11, color=INK, bold=False, after=6, line=1.25) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 14, 7
    else:
        size, color, before, after = 12, DARK_BLUE, 10, 5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, size=9, color=DARK_BLUE, bold=True, after=2, line=1.15)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, size=8.7, color=INK, after=2, line=1.15)
    set_table_fixed(table, widths)


def add_label_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Detail"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, size=9, color=DARK_BLUE, bold=True, after=2, line=1.15)
    for label, detail in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = detail
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for paragraph in row.cells[0].paragraphs:
            style_paragraph(paragraph, size=9.2, color=DARK_BLUE, bold=True, after=2, line=1.15)
        for paragraph in row.cells[1].paragraphs:
            style_paragraph(paragraph, size=9.2, color=INK, after=2, line=1.15)
    set_table_fixed(table, [1.181, 5.319])


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in (
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "STARCORE reference brief"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(header, size=9, color="6B7280", after=0, line=1)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("STARCORE Open Source Reference Roadmap")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.add_run("Verified 2026-06-03 · Paper 1.21.11 target · MIT-compatible implementation discipline")
    style_paragraph(subtitle, size=10.5, color="4B5563", after=10, line=1.15)

    add_heading(doc, "Executive Findings", 1)
    for item in [
        "The current STARCORE codebase already has a working strategic skeleton: module manager, services, nation, government, resolution, treasury, diplomacy, policy, map, war, technology, resource, officer, and event modules.",
        "Several older reference claims drifted: Towny latest is 0.103.0.0, not 0.103.1.0; Nations-Legacy GitHub did not resolve, while Spigot/Spiget reports v1.0.4.",
        "License hygiene matters: Towny is CC BY-NC-ND, WDE and Towns-and-Nations are GPL-3.0. Use them for concepts and architecture vocabulary only.",
        "The highest-value next step is not another dependency. It is a STARCORE-native repository layer plus a shared modifier/effect engine.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Local Baseline", 1)
    add_label_table(doc, [
        ("Compile target", "Java 21, Maven, Paper API 1.21.11-R0.1-SNAPSHOT."),
        ("Core wiring", "StarCorePlugin wires configuration, scheduler, event bus, persistence, internal permission/economy, ModuleManager, and services."),
        ("Persistence now", "Namespaced .properties files. Useful for V0.1, but not the final SQLite/MySQL data engine."),
        ("Tool evidence", "CodeGraph indexed 1,248 files / 29,234 nodes / 81,929 edges. Existing Graphify map AST captures the map module slice."),
    ])

    add_heading(doc, "External Reference Matrix", 1)
    add_table(
        doc,
        ["Reference", "Verified status", "License / risk", "STARCORE decision"],
        [
            ["Towny", "0.103.0.0, published 2026-05-08; MC 1.19-1.21 and 26.1.*.", "CC BY-NC-ND", "Study data decomposition, event surface, and persistence boundaries. No code copy."],
            ["Towns-and-Nations", "v1.0.0, published 2026-05-10; active RP town/nation plugin.", "GPL-3.0", "Study alliance/war UX and API boundary only."],
            ["NationTech-parent", "No releases; last pushed 2025-11-28.", "No declared license", "Idea reference for politics plus tech-tree coupling."],
            ["WorldDynamics Engine", "v0.2.1 release; v1.x Argon rewrite is unreleased.", "GPL-3.0", "Best conceptual match for legislature/economy/diplomacy. Avoid Towny/Vault coupling."],
            ["SkillTree / TechTree", "SkillTree v1.2.5 is old; TechTree 1.5.1 tested through MC 1.20.", "No clear license", "Study node/prerequisite/category vocabulary, not code."],
            ["Paper MenuType", "Official Paper docs live.", "N/A", "Preferred native in-game policy tree UI surface; verify exact API against Paper 1.21.11."],
            ["Okaeri Menu", "README dependency 2.0.1-beta.4; Paper 1.21+ and Java 21.", "MIT", "Strong UX reference for panes, async loading, and reactive states."],
            ["Inventory Framework", "v3.7.1, published 2025-11-30.", "MIT", "Fallback GUI library if MenuType cannot cover required UX."],
        ],
        [1.35, 1.75, 1.15, 2.25],
    )

    add_heading(doc, "Architecture Decisions", 1)
    for item in [
        "Keep STARCORE independent from Towny, Vault, LuckPerms, WorldDynamics, and PlaceholderAPI. Optional adapters can come later.",
        "Split repository responsibilities out of business modules before the SQLite/MySQL migration.",
        "Replace hard-coded policy/technology definitions with config-backed immutable runtime graphs.",
        "Add ModifierEngine as a shared effect surface for policy, government, diplomacy, treasury, war, technology, and resource systems.",
        "Use Paper MenuType as the first GUI prototype target; borrow Okaeri pane/reactive ideas without accepting a dependency by default.",
        "Turn the 300% performance goal into measurable budgets: claim lookup latency, policy activation latency, flush time, map render concurrency, and TPS impact under spark profiling. Treat Aikar flags as deployment baseline tuning, not a replacement for plugin optimization.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Supplemental Technical Decisions", 1)
    add_table(
        doc,
        ["Area", "Verified reference", "STARCORE decision"],
        [
            ["Module system", "JPMS docs live; PF4J release-3.15.0, Apache-2.0.", "Use Maven boundaries plus ServiceLoader/metadata first. Defer JPMS/PF4J until extension needs are proven."],
            ["Persistence", "Paper database guidance plus HikariCP.", "Use HikariCP-backed SQL repositories for SQLite/MySQL so default and large-server stores share connection discipline."],
            ["Config and text", "Configurate 4.2.0, Kyori Adventure 5.1.1.", "Move complex config to typed loaders and use Kyori Adventure/MiniMessage for player text."],
            ["Caching", "Caffeine v3.2.4, Apache-2.0.", "Use bounded caches for policy graph, relation snapshots, territory indexes, and map view models."],
            ["Geometry", "JTS 1.20.0; Poly2Tri Java forks are fragmented.", "Prefer JTS for containment/intersection. Keep Poly2Tri only for triangulation/visualization."],
            ["CI and profiling", "Minecraft Plugin Runtime Test is third-party; YourKit/JProfiler official pages are live.", "Use local tests and spark by default; reserve paid profilers for hard CPU/memory cases."],
        ],
        [1.15, 2.4, 2.95],
    )

    add_heading(doc, "V0.1 to V1 Implementation Path", 1)
    add_table(
        doc,
        ["Phase", "Implementation focus", "Acceptance signal"],
        [
            ["V0.1 hardening", "Repository interfaces; event classes; configurable policy defaults; module dependency tests.", "Maven tests pass; current .properties persistence remains wrapped behind repositories."],
            ["V0.5 core strategy", "ModifierEngine; multi-policy state; treasury ledger; treaty-backed diplomacy; MenuType prototype.", "Policy activation produces queryable effects and GUI state without direct domain mutation."],
            ["V1.0 launch", "SQLite default store; public Java API; CI matrix; spark budgets; Modrinth, Hangar, and SpigotMC release channels; optional compatibility adapters.", "Clean MIT release artifact plus measurable performance report on Paper target versions."],
        ],
        [1.15, 3.35, 2.0],
    )

    add_heading(doc, "Primary Sources", 1)
    for source in [
        "Towny: https://github.com/TownyAdvanced/Towny/releases/tag/0.103.0.0",
        "Towns-and-Nations: https://github.com/Leralix/Towns-and-Nations/releases/tag/v1.0.0",
        "WorldDynamics Engine: https://github.com/stoleyourharvs/WorldDynamics-Engine",
        "Paper MenuType API: https://docs.papermc.io/paper/dev/menu-type-api/",
        "Paper database guidance / HikariCP: https://docs.papermc.io/paper/dev/using-databases/ / https://github.com/brettwooldridge/HikariCP",
        "Okaeri Menu: https://github.com/OkaeriPoland/okaeri-menu",
        "Inventory Framework: https://github.com/devnatan/inventory-framework",
        "JPMS / PF4J: https://docs.oracle.com/javase/9/docs/api/java/lang/module/package-summary.html / https://github.com/pf4j/pf4j",
        "Configurate / Kyori Adventure: https://github.com/SpongePowered/Configurate / https://github.com/PaperMC/adventure",
        "Caffeine / JTS: https://github.com/ben-manes/caffeine / https://github.com/locationtech/jts",
        "Open plugin lists: https://www.spigotmc.org/wiki/list-of-open-source-plugins/ / https://github.com/LiteDevelopers/awesome-minecraft",
        "spark profiler: https://github.com/lucko/spark",
        "Aikar flags / flags.sh: https://docs.papermc.io/paper/aikars-flags/ / https://flags.sh/",
        "Release platforms: https://modrinth.com/ / https://hangar.papermc.io/ / https://www.spigotmc.org/resources/",
        "Paid profilers: https://www.yourkit.com/java/profiler/ / https://www.ej-technologies.com/products/jprofiler/overview.html",
    ]:
        add_bullet(doc, source)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
